"""
Utility functions for academic writing workflow
"""
from typing import Dict
import os
from datetime import datetime
from docx import Document

def save_to_docx(result: Dict, prompt: str) -> str:
    """Save the writing result to a DOCX file."""
    # Create filename based on prompt
    filename = f"{prompt[:50].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
    
    # Create document
    doc = Document()
    
    # Add content
    doc.add_heading(prompt, 0)
    
    # Add edited text if available, otherwise original text
    text = result["edited_text"] if result["edited_text"] else result["text"]
    doc.add_paragraph(text)
    
    # Save document
    doc.save(filename)
    return filename

def save_to_bib(result: Dict, prompt: str) -> str:
    """Save references to a BibTeX file."""
    if "references" not in result:
        return None
        
    # Create filename
    filename = f"{prompt[:50].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.bib"
    
    # Write references
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(result["references"])
        
    return filename
